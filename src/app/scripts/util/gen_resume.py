import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def add_heading_style(style, size, document):
    """
    Adds a custom heading style with a specified font size to the document.

    Args:
        style (str): The name of the style to add.
        size (int): The font size in points.
        document (Document): The Document object to which the style will be added.
    """
    heading_style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(size)
    heading_style.font.bold = True


def add_subheading_style(style, size, document):
    """
    Adds a custom subheading style with a specified font size to the document.

    Args:
        style (str): The name of the style to add.
        size (int): The font size in points.
        document (Document): The Document object to which the style will be added.
    """
    subheading_style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.size = Pt(size)
    subheading_style.font.bold = True


def generate_resume(resume_data):
    """
    Generates a resume document from the provided data.

    Args:
        resume_data (ResumeData): An object containing all necessary data to populate the resume.
    """
    document = Document()

    add_heading_style("CustomHeading1", 14, document)
    add_subheading_style("CustomHeading2", 12, document)

    document.add_heading(f"{resume_data.first_name} {resume_data.last_name}", level=1)
    document.add_heading("Contact Information", level=2)
    document.add_paragraph(f"Email: {resume_data.email}")
    document.add_paragraph(f"Phone: {resume_data.phone_number}")
    if resume_data.github_link:
        document.add_paragraph(f"GitHub: {resume_data.github_link}")
    if resume_data.linkedin_link:
        document.add_paragraph(f"LinkedIn: {resume_data.linkedin_link}")

    document.add_heading("Work Experience", level=2)
    for experience in resume_data.experiences:
        start_date = (
            experience.start_date.strftime("%B %Y")
            if experience.start_date
            else "Present"
        )
        end_date = (
            experience.end_date.strftime("%B %Y") if experience.end_date else "Present"
        )
        document.add_paragraph(
            f"{experience.title}, {experience.company_name}", style="CustomHeading2"
        )
        document.add_paragraph(f"{start_date} - {end_date}")

        for point in experience.bullet_points:
            document.add_paragraph(point, style="List Bullet")

    document.add_heading("Education", level=2)
    for education in resume_data.educations:
        document.add_paragraph(
            f"{education.school}, {education.major}", style="CustomHeading2"
        )
        if education.grad_year:
            document.add_paragraph(f"Graduated: {education.grad_year}")

    document.add_heading("Skills", level=2)
    for skill in resume_data.skills:
        document.add_paragraph(skill.name, style="List Bullet")

    document.add_heading("Projects", level=2)
    for project in resume_data.projects:
        document.add_paragraph(project.name, style="CustomHeading2")
        document.add_paragraph(f"Language Stack: {project.language_stack}")
        for point in project.bullet_points:
            document.add_paragraph(point, style="List Bullet")

    document.add_heading("Extracurricular Activities", level=2)
    for extracurricular in resume_data.extracurriculars:
        document.add_paragraph(extracurricular.name, style="CustomHeading2")
        for point in extracurricular.bullet_points:
            document.add_paragraph(point, style="List Bullet")

    document.save(os.getenv("SAVE_PATH"))
